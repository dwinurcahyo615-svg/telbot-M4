import pandas as pd
import os
import datetime
from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters
)
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from difflib import get_close_matches
from openai import OpenAI




# =========================
# TOKEN
# =========================
TOKEN = os.getenv("TOKEN")

# =========================
# Z.AI CONFIG
# =========================
ZAI_API_KEY = "e50ee8d0c03c44e9b44e527cc242929b.6CiBEiiIGdfFEZJe"

client_ai = OpenAI(
    api_key=ZAI_API_KEY,
    base_url="https://open.bigmodel.cn/api/paas/v4/"
)

# =========================
# GOOGLE SHEET CONFIG
# =========================
SPREADSHEET_ORDER_ID = "1ZeoREaTBeLFpwk2YH_UtAwpNDDIqoISpEcd8qBvg7VY"
GID_ORDER = "1672427230"

SPREADSHEET_EXCIS_ID = "1BVGqJwSrUGyuPlUFIbN9iG1tQ6OIyNkqubGkOV6BvdU"
GID_EXCIS = "387589510"

SPREADSHEET_MAPPING_ID = "1ch-wXzcOvyySMTFJubUNHz1FGrbEtvG6h132SoMlDeY"
GID_MAPPING = "1670929091"

SPREADSHEET_DTP_ID = "14Ekqu4eN5j33Y575VeWjMaZWcL_THssvIoic6cIm6Os"
GID_DTP = "2026604688"

sessions = {}

# ============================================================
# =================== DATABASE PELANGGAN =====================
# ============================================================
def load_dtp():

    url = f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_DTP_ID}/export?format=csv&gid={GID_DTP}"

    # baca tanpa header dulu
    df_raw = pd.read_csv(url, dtype=str, header=None)

    # cari baris header yang berisi PELANGGAN
    header_row = None

    for i in range(len(df_raw)):
        row = df_raw.iloc[i].fillna("").astype(str)
        row_text = " ".join(row).upper()

        if "PELANGGAN" in row_text and "EMAIL" in row_text:
            header_row = i
            break

    if header_row is None:
        raise Exception("Header DTP tidak ditemukan")

    # baca ulang dengan header yang benar
    df = pd.read_csv(url, dtype=str, header=header_row)

    # bersihkan nama kolom
    df.columns = df.columns.astype(str).str.strip().str.upper()

    # bersihkan isi
    for col in df.columns:
        df[col] = df[col].fillna("").astype(str).str.strip()

    return df

async def cekdtp(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not context.args:
        await update.message.reply_text(
            "Gunakan:\n/cekdtp NAMA PELANGGAN"
        )
        return

    keyword = " ".join(context.args).upper().strip()

    df = load_dtp()

    if "PELANGGAN" not in df.columns:
        await update.message.reply_text(
            f"Kolom PELANGGAN tidak ditemukan.\nKolom tersedia:\n{list(df.columns)}"
        )
        return

    hasil = df[df["PELANGGAN"].str.upper().str.contains(keyword, na=False)]

    if hasil.empty:
        await update.message.reply_text("PELANGGAN tidak ditemukan.")
        return

    # ============================
    # LOOP semua hasil
    # ============================

    pesan = f"🔎 Ditemukan {len(hasil)} data untuk: {keyword}\n\n"

    for i, row in hasil.iterrows():

        pelanggan = row.get("PELANGGAN", "-")
        id_pelanggan = row.get("ID PELANGGAN", "-")
        bw = row.get("BW", "-")
        email = row.get("EMAIL", "-")
        pic = row.get("PIC", "-")
        harga = row.get("HARGA", "-")

        pesan += (
            f"🏢 PELANGGAN: {pelanggan}\n"
            f"👤 PIC: {pic}\n"
            f"🔢 ID: {id_pelanggan}\n"
            f"🌐 BW: {bw}\n"
            f"📧 EMAIL: {email}\n"
            f"💰 HARGA: {harga}\n"
            f"{'-'*25}\n"
        )

    await update.message.reply_text(pesan)


# ============================================================
# ===================== MAPPING SECTION ======================
# ============================================================
async def mapping_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):

    await update.message.reply_text(

        "📊 MENU MAPPING AM\n\n"

        "/cek NAMA PELANGGAN\n"
        "/ceknm NAMA AM\n\n"

        "Contoh:\n"
        "/cek KOSPIN JASA\n"
        "/ceknm DANIEL"

    )

def load_mapping():

    url = f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_MAPPING_ID}/export?format=csv&gid={GID_MAPPING}"

    df = pd.read_csv(url)

    df.columns = df.columns.str.strip().str.upper()

    return df


def flexible_search(input_name, df):

    input_name = input_name.upper().strip()

    all_names = df["STANDARD NAME"].astype(str).str.upper()

    # exact contains
    contains = all_names[all_names.str.contains(input_name, na=False)]

    if not contains.empty:
        return contains.iloc[0]

    # reverse contains
    for name in all_names:
        if input_name in name or name in input_name:
            return name

    # fuzzy match
    match = get_close_matches(input_name, all_names.tolist(), n=1, cutoff=0.3)

    if match:
        return match[0]

    return None


def clean_ai_text(text):

    import re

    text = re.sub(r"\*\*", "", text)
    text = re.sub(r"\*", "", text)
    text = re.sub(r"`", "", text)

    return text.strip()

#CEK def
async def cek(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not context.args:
        await update.message.reply_text(
            "⚠️ Gunakan:\n/cek NAMA PELANGGAN"
        )
        return

    input_name = " ".join(context.args).upper()

    # tampilkan loading
    msg = await update.message.reply_text(
        "🔍 Sedang mencari data..."
    )

    df = load_mapping()

    df["STANDARD NAME"] = df["STANDARD NAME"].astype(str).str.upper()

    nama_fix = flexible_search(input_name, df)

    if not nama_fix:
        await msg.edit_text(
            "❌ STANDARD NAME tidak ditemukan."
        )
        return

    result = df[df["STANDARD NAME"] == nama_fix]

    context.user_data["last_standard_name"] = nama_fix
    context.user_data["full_dataframe"] = df

    nama_am = "-"
    revenue = "-"

    if "MAPPING NAMA AM 2026" in df.columns:

        vals = result["MAPPING NAMA AM 2026"].dropna().astype(str)
        vals = vals[vals.str.strip() != ""].unique().tolist()

        if vals:
            nama_am = "\n".join(vals)

    if "REVENUE 2025 (ALL BA)" in df.columns:

        vals = result["REVENUE 2025 (ALL BA)"].dropna().astype(str)
        vals = vals[vals.str.strip() != ""].unique().tolist()

        if vals:
            revenue = "\n".join(vals)

    pesan = (
        f"📌 STANDARD NAME: {nama_fix}\n\n"
        f"👤 MAPPING NAMA AM 2026:\n{nama_am}\n\n"
        f"💰 REVENUE 2025 (ALL BA):\n{revenue}\n\n"
        f"Ketik 'next' untuk melihat cabang lain di internet."
    )

    await msg.edit_text(pesan)
    
#CEK NAMA
async def ceknm(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not context.args:
        await update.message.reply_text("Gunakan:\n/ceknm NAMA AM")
        return

    keyword = " ".join(context.args).upper()

    df = load_mapping()

    df["NAMA AM"] = df["NAMA AM"].astype(str).str.upper()

    result = df[df["NAMA AM"].str.contains(keyword)]

    filename = f"MAPPING_{keyword}.xlsx"

    result.to_excel(filename, index=False)

    await update.message.reply_document(open(filename, "rb"))

    os.remove(filename)

# ============================================================
# ===================== ORDER SECTION ========================
# ============================================================
def load_sheet(spreadsheet_id, gid):
    url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv&gid={gid}"

    df_raw = pd.read_csv(url, dtype=str, header=None)

    header_row = None

    for i in range(len(df_raw)):
        row = df_raw.iloc[i].fillna("").astype(str)
        row_text = " ".join(row).upper()

        # cocok untuk ORDER
        if "NOMOR ORDER" in row_text:
            header_row = i
            break

        # cocok untuk EXCIS
        if "NAMA PIC TELKOM" in row_text:
            header_row = i
            break

    if header_row is None:
        raise Exception("Header tidak ditemukan")

    df = pd.read_csv(url, dtype=str, header=header_row)

    # bersihkan nama kolom
    df.columns = df.columns.astype(str).str.strip().str.upper()

    # ======================================
    # FIX DUPLICATE KET (khusus ORDER)
    # ======================================

    ket_indexes = [i for i, col in enumerate(df.columns) if col == "KET"]

    if len(ket_indexes) >= 2:
        df.columns.values[ket_indexes[0]] = "KET"
        df.columns.values[ket_indexes[1]] = "TGL PS"

    # ======================================

    # bersihkan isi
    for col in df.columns:
        df[col] = df[col].fillna("").astype(str).str.strip()

    return df



def format_order(row):
    return (
        f"📌 NOMOR ORDER: {row.get('NOMOR ORDER','-')}\n"
        f"📝 NOMOR INTERNET: {row.get('NOMOR INTERNET','-')}\n"
        f"📊 STATUS: {row.get('STATUS','-')}\n"
        f"🏬 TIPE: {row.get('TIPE ORDER DATIN','-')}\n"
        f"👤 NAMA AM: {row.get('NAMA AM','-')}\n"
        f"👤 PIC: {row.get('NAMA PIC','-')}\n"
        f"📞 NO HP PIC: {row.get('NO HP PIC','-')}\n"
        f"📅 TGL PS: {row.get('TGL PS','-')}\n"
        f"🏢 STO: {row.get('STO','-')}\n"
        f"🏢 NAMA PERUSAHAAN: {row.get('NAMA PERUSAHAAN','-')}\n"
        f"🏢 ALAMAT INSTALASI PELANGGAN: {row.get('ALAMAT INSTALASI PELANGGAN','-')}\n"
        f"📝 INPUT: {row.get('INPUT APA HARI INI?','-')}\n"
        f"📊 DEAL PAKET TLP ONLY: {row.get('DEAL PAKET TLP ONLY','-')}\n"
        f"📊 DEAL PAKET INDIBIZ 1:2: {row.get('DEAL PAKET INDIBIZ 1:2','-')}\n"
        f"📊 DEAL PAKET INDIBIZ 1:1: {row.get('DEAL PAKET INDIBIZ 1:1','-')}\n"
        f"📊 DEAL PAKET BUNDLING DIGI PRODUK: {row.get('DEAL PAKET BUNDLING DIGI PRODUK','-')}\n"
        f"📝 KET: {row.get('KET','-')}\n"
        
        
    )

async def order_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📌 MENU ORDER\n\n"
        "Cek berdasarkan Nomor Order:\n"
        "Contoh: /ord 1002186859\n\n"
        "Cek berdasarkan Nama AM:\n"
        "Contoh: /nama DANIEL:\n\n"
        "Cek berdasarkan Nomor Internet:\n"
        "Contoh: /net 142406119693"
    )


async def ord(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Contoh:\n/ord 1002186859")
        return

    nomor = context.args[0].strip()
    df = load_sheet(SPREADSHEET_ORDER_ID, GID_ORDER)

    hasil = df[df["NOMOR ORDER"] == nomor]

    if hasil.empty:
        await update.message.reply_text("Nomor order tidak ditemukan.")
        return

    await update.message.reply_text(format_order(hasil.iloc[0]))
    
async def net(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Contoh:\n/net 142406119693")
        return

    nomor_internet = context.args[0].strip()

    df = load_sheet(SPREADSHEET_ORDER_ID, GID_ORDER)

    hasil = df[df["NOMOR INTERNET"] == nomor_internet]

    if hasil.empty:
        await update.message.reply_text("Nomor internet tidak ditemukan.")
        return

    # 🔥 PENTING: pakai format_order biar sama seperti /ord
    await update.message.reply_text(format_order(hasil.iloc[0]))

async def nama(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Contoh:\n/nama DANIEL")
        return

    keyword = " ".join(context.args).strip().upper()
    df = load_sheet(SPREADSHEET_ORDER_ID, GID_ORDER)

    if "NAMA AM" not in df.columns:
        await update.message.reply_text("Kolom NAMA AM tidak ditemukan.")
        return

    hasil = df[df["NAMA AM"].str.upper().str.contains(keyword)]

    if hasil.empty:
        await update.message.reply_text("Nama tidak ditemukan.")
        return

    # ==============================
    # EXPORT KE EXCEL
    # ==============================
    filename = f"ORDER_{keyword}.xlsx"
    hasil.to_excel(filename, index=False)

    await update.message.reply_document(
        document=open(filename, "rb"),
        filename=filename,
        caption=f"Data Order untuk NAMA AM: {keyword}"
    )

    os.remove(filename)


async def excis(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not context.args:
        await update.message.reply_text("Contoh:\n/excis Mastur")
        return

    nama_pic = " ".join(context.args).upper()

    df = load_sheet(SPREADSHEET_EXCIS_ID, GID_EXCIS)

    kolom_pic = None

    for col in df.columns:
        if "PIC" in col.upper():
            kolom_pic = col
            break

    if kolom_pic is None:
        await update.message.reply_text(
            f"Kolom PIC tidak ditemukan.\nKolom tersedia:\n{list(df.columns)}"
        )
        return

    hasil = df[df[kolom_pic].astype(str).str.upper().str.contains(nama_pic, na=False)]

    if hasil.empty:
        await update.message.reply_text("Nama PIC tidak ditemukan.")
        return

    filename = f"DATA_{nama_pic}.xlsx"

    hasil.to_excel(filename, index=False)

    await update.message.reply_document(open(filename, "rb"))

    os.remove(filename)

# ============================================================
# ======================= SPH SECTION ========================
# ============================================================

def rupiah(n):
    return f"Rp {n:,}".replace(",", ".")

def replace_in_paragraph(paragraph, data):
    for run in paragraph.runs:
        for k, v in data.items():
            if k in run.text:
                run.text = run.text.replace(k, v)

async def buatsph(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    sessions[uid] = {"step": "nama", "items": []}
    await update.message.reply_text("Masukkan NAMA PERUSAHAAN:")


async def handle_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    text = update.message.text.strip()
    text_lower = text.lower()

# mapping
    if uid not in sessions and text_lower == "next":

        nama = context.user_data.get("last_standard_name")
        df = context.user_data.get("full_dataframe")

        if not nama:
            return

        msg = await update.message.reply_text(
            "🤖 Silahkan tunggu AI sedang mencari cabang lain..."
        )

        try:

            existing = df["STANDARD NAME"].astype(str).unique().tolist()
            existing_text = ", ".join(existing[:50])

            prompt = f"""
Cari SEMUA cabang lain dari perusahaan:
{nama}

Jangan sertakan cabang berikut:
{existing_text}

Sertakan:
- Nama cabang
- Alamat lengkap jika tersedia
"""

            response = client_ai.chat.completions.create(
                model="glm-4.5-flash",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            hasil = clean_ai_text(response.choices[0].message.content)

            await msg.edit_text(
                f"🔎 Cabang lain ditemukan AI:\n\n{hasil}"
            )

        except Exception as e:

            await msg.edit_text(
                "⚠️ Server AI sedang sibuk, coba lagi."
            )

            print(e)

        return

# sph
    if uid not in sessions:
        return

    s = sessions[uid]

    if s["step"] == "nama":
        s["nama"] = text
        s["step"] = "perihal"
        await update.message.reply_text("Masukkan PERIHAL:")

    elif s["step"] == "perihal":
        s["perihal"] = text
        s["step"] = "tanggal"
        await update.message.reply_text(
            "Masukkan TANGGAL (contoh: 10 FEBRUARI 2026)\nKetik next untuk hari ini"
        )

    elif s["step"] == "tanggal":
        if text.lower() == "next":
            s["tanggal"] = datetime.datetime.now().strftime("%d %B %Y").upper()
        else:
            s["tanggal"] = text.upper()
        s["step"] = "item_nama"
        await update.message.reply_text("Masukkan NAMA PRODUK:")

    elif s["step"] == "item_nama":
        s["current"] = {"item": text}
        s["step"] = "item_bandwidth"
        await update.message.reply_text("Masukkan BANDWIDTH:")

    elif s["step"] == "item_bandwidth":
        s["current"]["bandwidth"] = text
        s["step"] = "item_abonemen"
        await update.message.reply_text("Masukkan ABONEMEN / BLN (angka):")
    
    
    elif s["step"] == "item_abonemen":
        s["current"]["abonemen"] = int(text)
        s["step"] = "item_qty"
        await update.message.reply_text("Masukkan QTY:")
    
    
    elif s["step"] == "item_qty":
        s["current"]["qty"] = int(text)
        s["step"] = "item_psb"
        await update.message.reply_text("Masukkan BIAYA PSB (angka):")
    
    
    elif s["step"] == "item_psb":
        s["current"]["psb"] = int(text)
        s["step"] = "item_keterangan"
        await update.message.reply_text("Masukkan KETERANGAN:")

    elif s["step"] == "item_keterangan":
        s["current"]["keterangan"] = text
        s["items"].append(s["current"])
        s["step"] = "tambah"
        await update.message.reply_text(
            "Ketik 'tambah' untuk item lagi\nKetik 'selesai' jika sudah"
        )

    elif s["step"] == "tambah":
        if text.lower() == "tambah":
            s["step"] = "item_nama"
            await update.message.reply_text("Masukkan NAMA PRODUK:")
        else:
            s["step"] = "am_nama"
            await update.message.reply_text("Masukkan NAMA ACCOUNT MANAGER:")

    elif s["step"] == "am_nama":
        s["am_nama"] = text
        s["step"] = "am_jabatan"
        await update.message.reply_text(
            "Masukkan JABATAN AM\nKetik next untuk default (Account Manager)"
        )

    elif s["step"] == "am_jabatan":
        s["am_jabatan"] = "Account Manager" if text.lower() == "next" else text
        s["step"] = "alamat"
        await update.message.reply_text("Masukkan ALAMAT\nKetik next untuk default")

    elif s["step"] == "alamat":
        s["alamat"] = "JL PAHLAWAN NO.10 SEMARANG" if text.lower() == "next" else text
        s["step"] = "telepon"
        await update.message.reply_text("Masukkan TELEPON:")

    elif s["step"] == "telepon":
        s["telepon"] = text
        await generate_pdf(update, s)
        del sessions[uid]


async def generate_pdf(update, s):
    
    doc = Document("template.docx")  
    
    replace = {
        "{{NAMA_PERUSAHAAN}}": s["nama"],
        "{{PERIHAL}}": s["perihal"],
        "{{TANGGAL}}": s["tanggal"],
        "{{NAMA_AM}}": s["am_nama"],
        "{{JABATAN_AM}}": s["am_jabatan"],
        "{{ALAMAT}}": s["alamat"],
        "{{TELEPON}}": s["telepon"],
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, replace)

    # ================================
    # BAGIAN PEMBUATAN TABEL DINAMIS
    # ================================
    for p in doc.paragraphs:
        if "{{TABEL}}" in p.text:

            parent = p._element.getparent()
            index = parent.index(p._element)
            parent.remove(p._element)

            table = doc.add_table(rows=1, cols=7)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = False
            
            # =========================
            # AUTO FIT
            # =========================
            section = doc.sections[0]
            page_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
            
            table.columns[0].width = Inches(page_width * 0.05) #FIX
            table.columns[1].width = Inches(page_width * 0.13)
            table.columns[2].width = Inches(page_width * 0.08)
            table.columns[3].width = Inches(page_width * 0.06)
            table.columns[4].width = Inches(page_width * 0.16) #FIX
            table.columns[5].width = Inches(page_width * 0.12)
            table.columns[6].width = Inches(page_width * 0.13)

            def set_cell_color(cell, color):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), color)
                tcPr.append(shd)

            def format_cell(cell, bold=False, size=Pt(10)):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # center
                    for run in paragraph.runs:
                        run.font.name = "Poppins"
                        run.font.size = size
                        run.bold = bold

            # ROW 1 - NAMA PERUSAHAAN
            company_row = table.rows[0].cells
            merged_cell = company_row[0]

            for i in range(1, 7):
                merged_cell = merged_cell.merge(company_row[i])

            merged_cell.text = s["nama"]

            for paragraph in merged_cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)

            # HEADER
            header = table.add_row().cells
            headers = [
                            "NO",
                            "PRODUK",
                            "BW (MB)",
                            "QTY",
                            "ABONEMEN/BLN",
                            "BIAYA PSB",
                            "KET"
                        ]

            for i, text in enumerate(headers):
                header[i].text = text
                set_cell_color(header[i], "C00000")
                format_cell(header[i], bold=True, size=Pt(10))
                
                for paragraph in header[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # ITEM
            total_abon = 0
            total_psb = 0
            
            for idx, item in enumerate(s["items"], start=1):
            
                abon_total = item["abonemen"] * item["qty"]
                psb_total = item["psb"] * item["qty"]

                total_abon += abon_total
                total_psb += psb_total
            
                row = table.add_row().cells
            
                row[0].text = str(idx)
                row[1].text = item["item"]
                row[2].text = item["bandwidth"]
                row[3].text = str(item["qty"])
                row[4].text = rupiah(item["abonemen"])
                row[5].text = rupiah(item["psb"])
                row[6].text = item["keterangan"]





            for cell in row:
                format_cell(cell)

            #UBAH PPN DISINI
            ppn_abon = int(total_abon * 0.11)
            ppn_psb = int(total_psb * 0.11)
            
            sub_abon = total_abon + ppn_abon
            sub_psb = total_psb + ppn_psb

            grand_total = sub_abon + sub_psb

            #CENTER
            def center_bold(cell):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True

            # TOTAL
            row = table.add_row().cells

            label_cell = row[0]
            for i in range(1, 4):
                label_cell = label_cell.merge(row[i])

            label_cell.text = "TOTAL"

            row[4].text = rupiah(total_abon)
            row[5].text = rupiah(total_psb)

            # BOLD
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

            for i in [4,5]:
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            center_bold(label_cell)

            # PPN
            row = table.add_row().cells

            label_cell = row[0]
            for i in range(1, 4):
                label_cell = label_cell.merge(row[i])

            label_cell.text = "PPN"

            row[4].text = rupiah(ppn_abon)
            row[5].text = rupiah(ppn_psb)

            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

            for i in [4,5]:
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            center_bold(label_cell)
            
            # SUB TOTAL
            row = table.add_row().cells

            label_cell = row[0]
            for i in range(1, 4):
                label_cell = label_cell.merge(row[i])

            label_cell.text = "SUB TOTAL"

            # isi nilai
            row[4].text = rupiah(sub_abon)
            row[5].text = rupiah(sub_psb)

            # BOLD
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

            for i in [4,5]:
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            center_bold(label_cell)

            # GRAND TOTAL
            row = table.add_row().cells

            # kosongkan kiri
            label = row[0]
            for i in range(1, 4):
                label = label.merge(row[i])

            label.text = ""

            # 🔥 merge ABONEMEN + PSB (kolom 4 & 5)
            value = row[4]
            value = value.merge(row[5])

            value.text = rupiah(grand_total)

            # center + bold
            for paragraph in value.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

            # KET tetap ada
            row[6].text = ""

            parent.insert(index, table._tbl)
            # ================================
            # ATUR UKURAN FONT TABEL
            # ================================
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if i == 1:  # header row
                                run.font.name = "Arial"
                                run.font.size = Pt(10)

                            # wajib agar Word benar-benar pakai Poppins
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                            else:
                                run.font.name = "Arial"
                                run.font.size = Pt(10)

                            # wajib agar Word benar-benar pakai Poppins
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            break

    out_docx = f"SPH_{s['nama']}.docx"
    out_pdf = f"SPH_{s['nama']}.pdf"

    doc.save(out_docx)
    

    await update.message.reply_document(open(out_docx, "rb"))
    os.remove(out_docx)
    


# ============================================================
# ========================= START ============================
# ============================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🤖 BOT ALL IN ONE AKTIF\n\n"
        "/order → Menu Cek Order\n"
        "/excis NAMA → Exception Isolir\n"
        "/buatsph → Buat SPH\n"
        "/mapping → Mapping AM\n"
        "/cekdtp → Cek Data DTP\n"
    )

# ============================================================
# ========================= MAIN =============================
# ============================================================

def main():
    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("order", order_menu))
    app.add_handler(CommandHandler("ord", ord))
    app.add_handler(CommandHandler("nama", nama))
    app.add_handler(CommandHandler("net", net))
    app.add_handler(CommandHandler("excis", excis))
    app.add_handler(CommandHandler("buatsph", buatsph))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_all))
    app.add_handler(CommandHandler("cek", cek))
    app.add_handler(CommandHandler("ceknm", ceknm))
    app.add_handler(CommandHandler("mapping", mapping_menu))
    app.add_handler(CommandHandler("cekdtp", cekdtp))

    

    print("BOT ALL IN ONE BERJALAN...")
    app.run_polling()

if __name__ == "__main__":
    main()


























